# -*- coding: utf-8 -*-
"""
قطب‌نمای قرآنی — پردازش آینه‌ای (نسخهٔ موبایل / اندروید)
بازنویسی‌شده با Kivy — همهٔ قابلیت‌های نسخهٔ دسکتاپ، بدون وابستگی به PyQt.
روی ویندوز با پایتون اجرا می‌شود و با GitHub Actions به APK تبدیل می‌شود.
"""
import os
import json
import shutil
import zipfile
from datetime import datetime

from kivy.config import Config
Config.set('input', 'mouse', 'mouse,multitouch_on_demand')

from kivy.app import App
from kivy.core.text import LabelBase
from kivy.core.window import Window
from kivy.uix.screenmanager import ScreenManager, Screen, SlideTransition, FadeTransition
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.anchorlayout import AnchorLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.scrollview import ScrollView
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.image import Image
from kivy.uix.widget import Widget
from kivy.uix.popup import Popup
from kivy.uix.spinner import Spinner
from kivy.graphics import Color, RoundedRectangle, Rectangle
from kivy.animation import Animation
from kivy.clock import Clock
from kivy.metrics import dp
from kivy.utils import platform

import core
from rtl import rtl, rtl_multiline

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# اگر پوشهٔ assets وجود داشته باشد از آن استفاده می‌کنیم؛ وگرنه فایل‌ها را از همین پوشهٔ اصلی می‌خوانیم
# (روی گیت‌هاب فایل‌ها در ریشه هستند، روی ویندوز داخل assets)
ASSET_DIR = os.path.join(BASE_DIR, 'assets')
if not os.path.isdir(ASSET_DIR):
    ASSET_DIR = BASE_DIR


def asset(name):
    return os.path.join(ASSET_DIR, name)


# ---- ثبت فونت‌ها ----
LabelBase.register(name='arabic', fn_regular=asset('arabic.ttf'), fn_bold=asset('arabic_bold.ttf'))
LabelBase.register(name='ui', fn_regular=asset('font.ttf'), fn_bold=asset('font.ttf'))

# ---- پالت رنگ ----
C_BG = (0.05, 0.08, 0.14, 1)
C_PANEL = (1, 1, 1, 0.10)
C_PANEL_SOLID = (0.10, 0.14, 0.22, 1)
C_GOLD = (0.95, 0.77, 0.36, 1)
C_BLUE = (0.15, 0.55, 0.92, 1)
C_PURPLE = (0.61, 0.28, 0.80, 1)
C_ORANGE = (1.0, 0.60, 0.10, 1)
C_GREEN = (0.20, 0.72, 0.45, 1)
C_RED = (0.90, 0.28, 0.28, 1)
C_TEXT = (0.96, 0.97, 1, 1)
C_MUTED = (0.72, 0.78, 0.88, 1)


def P(text):
    """متن فارسی/عربی آمادهٔ نمایش."""
    return rtl_multiline(text)


# ==================================================================
# ویدجت‌های پایه
# ==================================================================
_MEASURE_CACHE = {}
_CORE_LABELS = {}


def _text_width(s, font_name, font_size):
    """عرض رندرشدهٔ یک رشته با همان فونت/اندازه."""
    key = (font_name, round(float(font_size), 1), s)
    v = _MEASURE_CACHE.get(key)
    if v is not None:
        return v
    try:
        from kivy.core.text import Label as CoreLabel
        ck = (font_name, round(float(font_size), 1))
        cl = _CORE_LABELS.get(ck)
        if cl is None:
            cl = CoreLabel(text='', font_name=font_name, font_size=font_size)
            _CORE_LABELS[ck] = cl
        cl.text = s
        cl.refresh()
        w = cl.content_width
    except Exception:
        w = len(s) * float(font_size) * 0.6
    if len(_MEASURE_CACHE) < 20000:
        _MEASURE_CACHE[key] = w
    return w


class RLabel(Label):
    """لیبل فارسی با شکل‌دهی راست‌به‌چپ و شکستن صحیح خطوط (رفع مشکل آینه‌ای)."""
    def __init__(self, text='', arabic=False, **kw):
        kw.setdefault('font_name', 'arabic' if arabic else 'ui')
        kw.setdefault('color', C_TEXT)
        kw.setdefault('halign', 'right')
        kw.setdefault('valign', 'middle')
        kw.setdefault('markup', False)
        self._raw = '' if text is None else str(text)
        super().__init__(**kw)
        self.bind(size=self._sync, font_size=self._sync)
        self._sync()

    def _wrap_para(self, para, max_w):
        words = para.split(' ')
        lines = []
        cur = ''
        for wd in words:
            trial = wd if not cur else cur + ' ' + wd
            if not cur or _text_width(rtl(trial), self.font_name, self.font_size) <= max_w:
                cur = trial
            else:
                lines.append(cur)
                cur = wd
        if cur:
            lines.append(cur)
        return lines

    def _sync(self, *a):
        w = self.width
        self.text_size = (w, None)
        raw = self._raw
        if not raw:
            self.text = ''
            return
        try:
            if w and w > 8:
                max_w = max(1.0, w - dp(6))
                out = []
                for para in raw.split('\n'):
                    if para == '':
                        out.append('')
                        continue
                    for ln in self._wrap_para(para, max_w):
                        out.append(rtl(ln))
                self.text = '\n'.join(out)
            else:
                self.text = P(raw)
        except Exception:
            self.text = P(raw)

    def set_text(self, text):
        self._raw = '' if text is None else str(text)
        self._sync()


class RoundBox(BoxLayout):
    """جعبهٔ گوشه‌گرد با پس‌زمینه."""
    def __init__(self, bg=C_PANEL_SOLID, radius=18, border=None, **kw):
        super().__init__(**kw)
        self._bg = bg
        self._radius = radius
        self._border = border
        with self.canvas.before:
            if border:
                self._bcol = Color(*border)
                self._brect = RoundedRectangle(radius=[radius])
            self._col = Color(*bg)
            self._rect = RoundedRectangle(radius=[radius])
        self.bind(pos=self._upd, size=self._upd)

    def _upd(self, *a):
        self._rect.pos = self.pos
        self._rect.size = self.size
        if self._border:
            self._brect.pos = (self.x - dp(1.5), self.y - dp(1.5))
            self._brect.size = (self.width + dp(3), self.height + dp(3))

    def set_bg(self, color):
        self._col.rgba = color


class PillButton(Button):
    """دکمهٔ گوشه‌گرد رنگی با انیمیشن فشردن."""
    def __init__(self, text='', bg=C_BLUE, fg=(1, 1, 1, 1), radius=14, font_size='16sp', **kw):
        super().__init__(**kw)
        self.background_normal = ''
        self.background_color = (0, 0, 0, 0)
        self.font_name = 'ui'
        self.color = fg
        self.font_size = font_size
        self.text = P(text)
        self._bg = list(bg)
        self._radius = radius
        with self.canvas.before:
            self._col = Color(*self._bg)
            self._rect = RoundedRectangle(radius=[radius])
        self.bind(pos=self._upd, size=self._upd, state=self._state)

    def _upd(self, *a):
        self._rect.pos = self.pos
        self._rect.size = self.size

    def _state(self, *a):
        if self.state == 'down':
            self._col.rgba = [min(1, c * 1.25) for c in self._bg[:3]] + [self._bg[3]]
        else:
            self._col.rgba = self._bg

    def set_text(self, text):
        self.text = P(text)


class PersianTextInput(TextInput):
    # فیلد متنی راست‌به‌چپ که هنگام تایپ، فارسی را درست (بدون آینه‌ای) نشان می‌دهد
    def __init__(self, on_change=None, **kw):
        kw.setdefault('font_name', 'ui')
        kw.setdefault('halign', 'right')
        kw.setdefault('multiline', False)
        super().__init__(**kw)
        self._logical = ''
        self._guard = False
        self.on_change = on_change

    def _render(self):
        self._guard = True
        try:
            self.text = rtl(self._logical) if self._logical else ''
            self.cursor = (len(self.text), 0)
        finally:
            self._guard = False
        if self.on_change:
            self.on_change(self._logical)

    def insert_text(self, substring, from_undo=False):
        if self._guard or from_undo:
            return super().insert_text(substring, from_undo=from_undo)
        self._logical += substring
        self._render()
        return None

    def do_backspace(self, from_undo=False, mode='bkspc'):
        if self._guard or from_undo:
            return super().do_backspace(from_undo=from_undo, mode=mode)
        if self._logical:
            self._logical = self._logical[:-1]
            self._render()
        return None

    @property
    def query(self):
        return self._logical

    def set_logical(self, value):
        self._logical = value or ''
        self._render()

    def clear_logical(self):
        self._logical = ''
        self._guard = True
        try:
            self.text = ''
        finally:
            self._guard = False


def toast(message, title='پیام'):
    content = BoxLayout(orientation='vertical', padding=dp(16), spacing=dp(12))
    content.add_widget(RLabel(message, font_size='16sp', halign='center'))
    p = Popup(title=P(title), content=content, size_hint=(0.85, 0.4),
              title_font='ui', title_align='center', separator_color=C_GOLD)
    btn = PillButton('باشه', bg=C_BLUE, size_hint_y=None, height=dp(46))
    btn.bind(on_release=p.dismiss)
    content.add_widget(btn)
    p.open()
    return p


def confirm(message, on_yes, title='تأیید'):
    content = BoxLayout(orientation='vertical', padding=dp(16), spacing=dp(12))
    content.add_widget(RLabel(message, font_size='16sp', halign='center'))
    row = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(10))
    p = Popup(title=P(title), content=content, size_hint=(0.85, 0.42),
              title_font='ui', title_align='center', separator_color=C_GOLD)
    yes = PillButton('بله', bg=C_GREEN)
    no = PillButton('انصراف', bg=C_RED)
    def _yes(*a):
        p.dismiss()
        on_yes()
    yes.bind(on_release=_yes)
    no.bind(on_release=p.dismiss)
    row.add_widget(yes)
    row.add_widget(no)
    content.add_widget(row)
    p.open()
    return p


# ==================================================================
# صفحهٔ پایه (پس‌زمینه + هدر)
# ==================================================================
class BaseScreen(Screen):
    def __init__(self, title='', show_back=True, **kw):
        super().__init__(**kw)
        self.root_layout = FloatLayout()
        add = self.root_layout.add_widget
        # پس‌زمینه
        with self.root_layout.canvas.before:
            Color(*C_BG)
            self._bgrect = Rectangle(pos=(0, 0), size=Window.size)
        Window.bind(size=lambda *a: setattr(self._bgrect, 'size', Window.size))
        try:
            self.bg_image = Image(source=asset('bg.jpg'), allow_stretch=True,
                                  keep_ratio=False, color=(1, 1, 1, 0.28),
                                  size_hint=(1, 1), pos_hint={'x': 0, 'y': 0})
            add(self.bg_image)
        except Exception:
            pass
        # ستون اصلی
        self.container = BoxLayout(orientation='vertical', size_hint=(1, 1),
                                   padding=dp(10), spacing=dp(8))
        add(self.container)
        # هدر
        header = RoundBox(bg=(1, 1, 1, 0.06), orientation='horizontal',
                          size_hint_y=None, height=dp(56), padding=dp(8), spacing=dp(6))
        if show_back:
            back = PillButton('بازگشت', bg=(1, 1, 1, 0.14), size_hint_x=None, width=dp(110),
                              font_size='14sp')
            back.bind(on_release=self.go_back)
            header.add_widget(back)
        else:
            header.add_widget(Widget(size_hint_x=None, width=dp(4)))
        self.title_label = RLabel(title, font_name='ui', bold=True, font_size='19sp',
                                  halign='center', color=C_GOLD)
        header.add_widget(self.title_label)
        header.add_widget(Widget(size_hint_x=None, width=dp(110) if show_back else dp(4)))
        self.container.add_widget(header)
        self.add_widget(self.root_layout)

    def go_back(self, *a):
        self.manager.transition = SlideTransition(direction='right')
        self.manager.current = 'home'

    def body(self, widget):
        self.container.add_widget(widget)


# ==================================================================
# خانه
# ==================================================================
class HomeScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='', show_back=False, **kw)
        app = App.get_running_app()
        scroll = ScrollView(size_hint=(1, 1), do_scroll_x=False, bar_width=dp(4))
        content = BoxLayout(orientation='vertical', size_hint_y=None, padding=dp(6), spacing=dp(12))
        content.bind(minimum_height=content.setter('height'))

        # آیهٔ محوری متحرک
        self.verse = RLabel('« إِنَّا نَحْنُ نَزَّلْنَا الذِّکْرَ وَإِنَّا لَهُ لَحَافِظُونَ »',
                             arabic=True, font_size='21sp', halign='center', color=C_GOLD,
                             size_hint_y=None, height=dp(64))
        content.add_widget(self.verse)
        anim = (Animation(opacity=0.35, duration=1.8) + Animation(opacity=1, duration=1.8))
        anim.repeat = True
        anim.start(self.verse)

        title = RLabel('قطب‌نمای قرآنی', bold=True, font_size='30sp', halign='center',
                       color=C_TEXT, size_hint_y=None, height=dp(46))
        subtitle = RLabel('پردازش آینه‌ای (هولوگرافیک)', font_size='15sp', halign='center',
                          color=C_MUTED, size_hint_y=None, height=dp(28))
        content.add_widget(title)
        content.add_widget(subtitle)

        # پنل ورودی بذر
        seedbox = RoundBox(bg=(1, 1, 1, 0.09), orientation='vertical', size_hint_y=None,
                           height=dp(230), padding=dp(14), spacing=dp(10))
        seedbox.add_widget(RLabel('انتخاب بذر (سوره و آیه)', bold=True, font_size='17sp',
                                  halign='center', color=C_GOLD, size_hint_y=None, height=dp(28)))
        inrow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(10))
        self.in_s = TextInput(hint_text=P('سوره'), multiline=False, font_name='ui',
                              halign='center', font_size='18sp', input_filter='int',
                              background_color=(1, 1, 1, 0.92), foreground_color=(0.05, 0.08, 0.14, 1),
                              padding=[dp(8), dp(12)])
        self.in_a = TextInput(hint_text=P('آیه'), multiline=False, font_name='ui',
                              halign='center', font_size='18sp', input_filter='int',
                              background_color=(1, 1, 1, 0.92), foreground_color=(0.05, 0.08, 0.14, 1),
                              padding=[dp(8), dp(12)])
        inrow.add_widget(self.in_s)
        inrow.add_widget(self.in_a)
        seedbox.add_widget(inrow)

        brow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        b_matrix = PillButton('پردازش ماتریس', bg=C_BLUE)
        b_matrix.bind(on_release=lambda *a: self.run('matrix'))
        brow.add_widget(b_matrix)
        seedbox.add_widget(brow)

        prow = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        b_sem = PillButton('پیش‌بینی (معنا)', bg=C_ORANGE)
        b_sem.bind(on_release=lambda *a: self.run('semantic'))
        b_num = PillButton('پیش‌بینی (اعداد)', bg=C_PURPLE)
        b_num.bind(on_release=lambda *a: self.run('numeric'))
        prow.add_widget(b_sem)
        prow.add_widget(b_num)
        seedbox.add_widget(prow)
        content.add_widget(seedbox)

        # کاشی‌های ناوبری
        grid = GridLayout(cols=2, size_hint_y=None, spacing=dp(10))
        grid.bind(minimum_height=grid.setter('height'))
        nav = [
            ('لابراتوار کشفیات', C_GREEN, 'lab'),
            ('گلچین برگزیده', C_GOLD, 'featured'),
            ('جستجوی کشفیات', C_BLUE, 'search'),
            ('مدیریت برچسب‌ها', C_PURPLE, 'tags'),
            ('رسانه و معرفی', C_ORANGE, 'media'),
            ('راهنما', (0.3, 0.4, 0.55, 1), 'guide'),
            ('پشتیبان و بازیابی', (0.25, 0.5, 0.6, 1), 'backup'),
            ('درباره', (0.4, 0.35, 0.5, 1), 'about'),
        ]
        for label, color, target in nav:
            b = PillButton(label, bg=color, size_hint_y=None, height=dp(64), font_size='15sp')
            b.bind(on_release=lambda inst, t=target: self.nav(t))
            grid.add_widget(b)
        content.add_widget(grid)
        content.add_widget(Widget(size_hint_y=None, height=dp(20)))

        scroll.add_widget(content)
        self.body(scroll)

    def _seed(self):
        try:
            s = int(core.conv(self.in_s.text.strip()))
            a = int(core.conv(self.in_a.text.strip()))
            return s, a
        except Exception:
            toast('لطفاً شمارهٔ سوره و آیهٔ معتبر وارد کنید.', 'خطا')
            return None

    def run(self, kind):
        seed = self._seed()
        if not seed:
            return
        app = App.get_running_app()
        s, a = seed
        if kind == 'matrix':
            fs, fa, is_fb, msg = app.data.apply_circular(s, a)
            if fs is None:
                toast('آیهٔ معتبر یافت نشد.', 'خطا')
                return
            scr = self.manager.get_screen('matrix')
            scr.show(fs, fa)
            self.manager.transition = SlideTransition(direction='left')
            self.manager.current = 'matrix'
        else:
            found = app.data.find_seed(s, a)
            if not found:
                toast('آیهٔ مورد نظر در دیتابیس یافت نشد.', 'خطا')
                return
            fs, fa = found
            scr = self.manager.get_screen('predict')
            scr.show(fs, fa, kind)
            self.manager.transition = SlideTransition(direction='left')
            self.manager.current = 'predict'

    def nav(self, target):
        scr = self.manager.get_screen(target)
        if hasattr(scr, 'refresh'):
            scr.refresh()
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = target


# ==================================================================
# کارت آیه (مشترک)
# ==================================================================
def verse_card(mode, s, a, arb, pers, is_seed=False, is_fallback=False, reason='',
               on_save=None, score_text=''):
    bg = (0.16, 0.13, 0.05, 1) if is_seed else ((0.22, 0.08, 0.08, 1) if is_fallback else (0.10, 0.14, 0.22, 1))
    border = C_GOLD if is_seed else (C_RED if is_fallback else None)
    card = RoundBox(bg=bg, border=border, orientation='vertical', size_hint_y=None,
                    padding=dp(12), spacing=dp(6))
    head = RLabel(f'{mode}   سوره {s} ، آیه {a}', bold=True, font_size='15sp',
                  color=(C_GOLD if is_seed else C_ORANGE), halign='right', size_hint_y=None)
    head.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
    card.add_widget(head)
    if score_text:
        sc = RLabel(score_text, font_size='13sp', color=C_MUTED, halign='right', size_hint_y=None)
        sc.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(2)))
        card.add_widget(sc)
    arb_l = RLabel(f'« {arb} »', arabic=True, font_size='20sp', halign='center',
                   color=C_TEXT, size_hint_y=None)
    arb_l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
    card.add_widget(arb_l)
    pers_l = RLabel(pers, font_size='14sp', halign='center', color=C_MUTED, size_hint_y=None)
    pers_l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
    card.add_widget(pers_l)
    if is_fallback and reason:
        warn = RLabel('' + reason, font_size='12sp', halign='right', color=C_RED, size_hint_y=None)
        warn.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(warn)
    if on_save:
        btn = PillButton('ثبت این کشف', bg=C_GREEN, size_hint_y=None, height=dp(42), font_size='14sp')
        btn.bind(on_release=lambda *x: on_save())
        card.add_widget(btn)

    def _h(*a):
        total = sum(c.height for c in card.children) + dp(24) + dp(6) * (len(card.children) - 1)
        card.height = total
    Clock.schedule_once(_h, 0)
    card.bind(minimum_height=lambda i, v: setattr(card, 'height', v))
    return card


# ==================================================================
# صفحهٔ ماتریس
# ==================================================================
class MatrixScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='پردازش ماتریس آینه‌ای', **kw)
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def show(self, s, a):
        app = App.get_running_app()
        self.list.clear_widgets()
        cards = core.process_matrix(app.data, s, a)
        for i, c in enumerate(cards):
            on_save = None
            seed = cards[0]
            if c['kind'] == 'target':
                on_save = (lambda cc=c, sd=seed: app.add_discovery(sd, cc))
            w = verse_card(c['mode'], c['s'], c['a'], c['arb'], c['pers'],
                           is_seed=(c['kind'] == 'seed'), is_fallback=c['is_fallback'],
                           reason=c['reason'], on_save=on_save)
            w.opacity = 0
            self.list.add_widget(w)
            Animation(opacity=1, duration=0.35).start(w)


# ==================================================================
# صفحهٔ پیش‌بینی (معنا / اعداد)
# ==================================================================
class PredictScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='پیش‌بینی آینه', **kw)
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def show(self, s, a, kind):
        app = App.get_running_app()
        self.list.clear_widgets()
        seed = app.data.get(s, a)
        self.title_label.set_text('پیش‌بینی ' + ('معنایی' if kind == 'semantic' else 'عددی'))
        # کارت بذر
        self.list.add_widget(verse_card('بذر ساختاری', s, a, seed['arb'], seed['pers'], is_seed=True))
        if kind == 'semantic':
            preds = core.predict_mirror(app.data, s, a, seed['arb'])
            for rank, (op, ts, ta, score, status, is_fb, msg) in enumerate(preds, 1):
                d = app.data.get(ts, ta) or {}
                sc = f'رتبه {rank} | امتیاز: {score:.0f}٪ | {status}'
                seed_card = {'mode': op, 's': s, 'a': a, 'arb': seed['arb'], 'pers': seed['pers']}
                tgt = {'mode': op, 's': ts, 'a': ta, 'arb': d.get('arb', ''), 'pers': d.get('pers', ''),
                       'is_fallback': is_fb, 'reason': msg}
                w = verse_card(op, ts, ta, d.get('arb', ''), d.get('pers', ''),
                               is_fallback=is_fb, reason=(msg if is_fb else ''),
                               score_text=sc,
                               on_save=(lambda sd=seed_card, cc=tgt: app.add_discovery(sd, cc)))
                w.opacity = 0
                self.list.add_widget(w)
                Animation(opacity=1, duration=0.3).start(w)
        else:
            preds = core.predict_mirror_numeric(app.data, s, a)
            if not preds:
                self.list.add_widget(RLabel('هیچ مقصد معتبری با الگوریتم عددی یافت نشد.',
                                            font_size='15sp', halign='center', color=C_MUTED,
                                            size_hint_y=None, height=dp(60)))
                return
            for (op, ts, ta, prio, detail, is_fb, msg) in preds:
                d = app.data.get(ts, ta) or {}
                sc = f'اولویت {prio} | {detail}'
                seed_card = {'mode': op, 's': s, 'a': a, 'arb': seed['arb'], 'pers': seed['pers']}
                tgt = {'mode': op, 's': ts, 'a': ta, 'arb': d.get('arb', ''), 'pers': d.get('pers', ''),
                       'is_fallback': is_fb, 'reason': msg}
                w = verse_card(op, ts, ta, d.get('arb', ''), d.get('pers', ''),
                               is_fallback=is_fb, reason=(msg if is_fb else ''), score_text=sc,
                               on_save=(lambda sd=seed_card, cc=tgt: app.add_discovery(sd, cc)))
                w.opacity = 0
                self.list.add_widget(w)
                Animation(opacity=1, duration=0.3).start(w)


# ==================================================================
# لابراتوار کشفیات
# ==================================================================
OPERATORS = [
    ('T1', 'جابجایی خالص بذر'),
    ('T2', 'تقارن درجا کامل'),
    ('T3', 'تقارن درجا فقط سوره'),
    ('T4', 'تقارن درجا فقط آیه'),
    ('T5', 'جابجایی + تقارن کامل'),
    ('T6', 'جابجایی + تقارن فقط سوره'),
    ('T7', 'جابجایی + تقارن فقط آیه'),
    ('OTHER', 'گروهی و سایر'),
]


def op_of(item):
    m = str(item.get('mode', ''))
    if 'خالص' in m:
        return 'T1'
    if 'ضربدری' in m:
        return 'T5'
    if 'جابجایی' in m and 'کامل' in m:
        return 'T5'
    if 'تقارن درجا کامل' in m:
        return 'T2'
    if 'جابجایی' in m and 'فقط سوره' in m:
        return 'T6'
    if 'جابجایی' in m and 'فقط آیه' in m:
        return 'T7'
    if 'فقط سوره' in m:
        return 'T3'
    if 'فقط آیه' in m:
        return 'T4'
    return 'OTHER'


def _auto_label(text, arabic=False, **kw):
    kw.setdefault('size_hint_y', None)
    lbl = RLabel(text, arabic=arabic, **kw)
    lbl.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(8)))
    return lbl


def _verse_block(border, s, a, arb, pers):
    c = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=border, orientation='vertical',
                 size_hint_y=None, padding=dp(10), spacing=dp(4))
    c.add_widget(_auto_label('سوره %s ، آیه %s' % (s, a), font_size='12sp', color=C_MUTED, halign='right'))
    c.add_widget(_auto_label('« %s »' % (arb or ''), arabic=True, font_size='18sp', color=C_TEXT, halign='center'))
    c.add_widget(_auto_label('ترجمه: ' + (pers or ''), font_size='13sp', color=C_MUTED, halign='right'))
    c.bind(minimum_height=lambda i, v: setattr(c, 'height', v + dp(24)))
    return c


def show_discovery(item, source='lab', screen=None):
    app = App.get_running_app()
    root = BoxLayout(orientation='vertical', padding=dp(8), spacing=dp(8))
    scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
    box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
    box.bind(minimum_height=box.setter('height'))
    box.add_widget(_auto_label('آیهٔ مبدأ (بذر)', bold=True, font_size='15sp', color=C_GOLD, halign='right'))
    box.add_widget(_verse_block(C_PURPLE, item.get('seed_s'), item.get('seed_a'),
                                item.get('seed_arb', ''), item.get('seed_pers', '')))
    box.add_widget(_auto_label('گرهٔ کشف‌شده: ' + str(item.get('mode', '')), bold=True,
                               font_size='14sp', color=C_GOLD, halign='right'))
    box.add_widget(_verse_block(C_GOLD, item.get('target_s'), item.get('target_a'),
                                item.get('target_arb', ''), item.get('target_pers', '')))
    box.add_widget(_auto_label('رفتار شبکه: ' + str(item.get('relation_type', 'نامشخص')),
                               font_size='14sp', color=C_ORANGE, halign='right'))
    _note = item.get('note', '')
    box.add_widget(_auto_label('تحلیل شما: ' + (_note if _note else '—'),
                               font_size='14sp', color=C_TEXT, halign='right'))
    scroll.add_widget(box)
    root.add_widget(scroll)
    p = Popup(title=P('جزئیات کشف'), content=root, size_hint=(0.96, 0.9),
              title_font='ui', title_align='center', separator_color=C_GOLD)

    def _refresh_parent():
        if screen is not None and hasattr(screen, 'refresh'):
            screen.refresh()

    def _copy(*a):
        try:
            from kivy.core.clipboard import Clipboard
            txt = ('[%s] رفتار: %s\n' % (item.get('mode', ''), item.get('relation_type', ''))
                   + 'مبدأ (سوره %s آیه %s): %s\n%s\n' % (item.get('seed_s'), item.get('seed_a'),
                     item.get('seed_arb', ''), item.get('seed_pers', ''))
                   + 'مقصد (سوره %s آیه %s): %s\n%s\n' % (item.get('target_s'), item.get('target_a'),
                     item.get('target_arb', ''), item.get('target_pers', ''))
                   + ('تحلیل: ' + _note if _note else ''))
            Clipboard.copy(txt)
            toast('اطلاعات کشف کپی شد.', 'کپی')
        except Exception:
            toast('کپی ممکن نشد.', 'خطا')

    def _delete(*a):
        def _do():
            lst = app.favs if source == 'lab' else app.featured
            key = (item.get('seed_s'), item.get('seed_a'), item.get('target_s'),
                   item.get('target_a'), item.get('mode'))
            for i, it in enumerate(lst):
                if (it.get('seed_s'), it.get('seed_a'), it.get('target_s'),
                        it.get('target_a'), it.get('mode')) == key:
                    del lst[i]
                    break
            if source == 'lab':
                app.save_favs()
            else:
                app.save_featured()
            p.dismiss()
            _refresh_parent()
        confirm('این کشف حذف شود؟', _do, 'حذف کشف')

    def _to_featured(*a):
        app.add_featured(item)

    def _edit(*a):
        content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(10))
        content.add_widget(RLabel('تحلیل شما:', font_size='15sp', size_hint_y=None, height=dp(26)))
        ti = PersianTextInput(multiline=True, font_size='15sp',
                              size_hint_y=None, height=dp(120), background_color=(1, 1, 1, 0.95),
                              foreground_color=(0.05, 0.08, 0.14, 1))
        ti.set_logical(item.get('note', ''))
        content.add_widget(ti)
        tags = app.get_all_tags()
        content.add_widget(RLabel('برچسب (رفتار شبکه):', font_size='15sp', size_hint_y=None, height=dp(26)))
        sp = Spinner(text=P(item.get('relation_type', 'نامشخص')), values=[P(t) for t in tags],
                     font_name='ui', size_hint_y=None, height=dp(44))
        content.add_widget(sp)
        tag_map = {P(t): t for t in tags}
        ep = Popup(title=P('ویرایش تحلیل'), content=content, size_hint=(0.92, 0.6),
                   title_font='ui', title_align='center', separator_color=C_GOLD)
        row2 = BoxLayout(size_hint_y=None, height=dp(46), spacing=dp(10))
        sv = PillButton('ذخیره', bg=C_GREEN)

        def _sv(*a):
            item['note'] = ti.query
            item['relation_type'] = tag_map.get(sp.text, 'نامشخص')
            if source == 'lab':
                app.save_favs()
            else:
                app.save_featured()
            ep.dismiss()
            p.dismiss()
            _refresh_parent()
        sv.bind(on_release=_sv)
        cn = PillButton('انصراف', bg=C_RED)
        cn.bind(on_release=ep.dismiss)
        row2.add_widget(sv)
        row2.add_widget(cn)
        content.add_widget(row2)
        ep.open()

    grid = GridLayout(cols=2, size_hint_y=None, height=dp(104), spacing=dp(8))
    if source == 'lab':
        bb = PillButton('افزودن به گلچین', bg=C_GOLD, font_size='14sp')
        bb.bind(on_release=_to_featured)
        grid.add_widget(bb)
    be = PillButton('ویرایش تحلیل', bg=C_BLUE, font_size='14sp')
    be.bind(on_release=_edit)
    grid.add_widget(be)
    bd = PillButton('حذف کشف' if source == 'lab' else 'حذف از گلچین', bg=C_RED, font_size='14sp')
    bd.bind(on_release=_delete)
    grid.add_widget(bd)
    bc = PillButton('کپی اطلاعات', bg=C_GREEN, font_size='14sp')
    bc.bind(on_release=_copy)
    grid.add_widget(bc)
    root.add_widget(grid)
    close = PillButton('بستن', bg=(1, 1, 1, 0.14), size_hint_y=None, height=dp(46))
    close.bind(on_release=p.dismiss)
    root.add_widget(close)
    p.open()


class LabScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='لابراتوار کشفیات', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(8))
        b_all = PillButton('افزودن همه به گلچین', bg=C_GOLD, font_size='14sp')
        b_all.bind(on_release=lambda *a: self.add_all_featured())
        top.add_widget(b_all)
        self.count_lbl = RLabel('', font_size='14sp', halign='center', color=C_MUTED)
        top.add_widget(self.count_lbl)
        self.body(top)
        self.body(RLabel('برای دیدن کشف‌های هر عملگر، روی آن بزنید.', font_size='13sp',
                         halign='center', color=C_MUTED, size_hint_y=None, height=dp(26)))
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.grid = GridLayout(cols=1, size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.grid.bind(minimum_height=self.grid.setter('height'))
        self.scroll.add_widget(self.grid)
        self.body(self.scroll)

    def refresh(self):
        app = App.get_running_app()
        self.grid.clear_widgets()
        self.count_lbl.set_text('%d کشف' % len(app.favs))
        if not app.favs:
            self.grid.add_widget(RLabel('هنوز کشفی ثبت نشده است. از صفحهٔ پردازش، کشف ثبت کنید.',
                                        font_size='15sp', halign='center', color=C_MUTED,
                                        size_hint_y=None, height=dp(80)))
            return
        counts = {}
        for it in app.favs:
            k = op_of(it)
            counts[k] = counts.get(k, 0) + 1
        for key, title in OPERATORS:
            n = counts.get(key, 0)
            if n == 0:
                continue
            b = PillButton('%s  (%d کشف)' % (title, n), bg=(0.16, 0.13, 0.05, 1),
                           size_hint_y=None, height=dp(64), font_size='16sp')
            b.bind(on_release=lambda inst, k=key, t=title: self.open_op(k, t))
            self.grid.add_widget(b)

    def open_op(self, key, title):
        scr = self.manager.get_screen('operator')
        scr.load('lab', key, title)
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = 'operator'

    def add_all_featured(self):
        app = App.get_running_app()
        n = app.add_all_featured()
        toast('%d کشف به گلچین اضافه شد.' % n, 'گلچین')


class OperatorScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='کشف‌های عملگر', **kw)
        self.source = 'lab'
        self.op_key = 'T1'
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def load(self, source, op_key, title):
        self.source = source
        self.op_key = op_key
        src_name = 'لابراتوار' if source == 'lab' else 'گلچین'
        self.title_label.set_text('%s — %s' % (src_name, title))
        self.refresh()

    def refresh(self):
        app = App.get_running_app()
        self.list.clear_widgets()
        items = app.favs if self.source == 'lab' else app.featured
        matched = [it for it in items if op_of(it) == self.op_key]
        if not matched:
            self.list.add_widget(RLabel('کشفی در این عملگر نیست.', font_size='15sp',
                                        halign='center', color=C_MUTED, size_hint_y=None, height=dp(70)))
            return
        for it in matched:
            self.list.add_widget(self._card(it))

    def _card(self, item):
        border = C_GOLD if self.source == 'featured' else C_BLUE
        card = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=border, orientation='vertical',
                        size_hint_y=None, padding=dp(10), spacing=dp(4))
        pair = RLabel('سوره %s:%s     سوره %s:%s' % (item.get('seed_s'), item.get('seed_a'),
                      item.get('target_s'), item.get('target_a')),
                      font_size='13sp', color=C_MUTED, halign='center', size_hint_y=None)
        pair.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(pair)
        a2 = RLabel('« %s »' % (item.get('target_arb', '')), arabic=True, font_size='16sp',
                    halign='center', color=C_TEXT, size_hint_y=None)
        a2.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        card.add_widget(a2)
        rel = RLabel('رفتار: %s' % item.get('relation_type', 'نامشخص'), font_size='12sp',
                     color=C_ORANGE, halign='right', size_hint_y=None)
        rel.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(2)))
        card.add_widget(rel)
        btn = PillButton('مشاهدهٔ جزئیات', bg=C_BLUE, size_hint_y=None, height=dp(40), font_size='13sp')
        btn.bind(on_release=lambda *a, it=item: show_discovery(it, self.source, self))
        card.add_widget(btn)
        card.bind(minimum_height=lambda i, v: setattr(card, 'height', v + dp(20)))
        return card

    def go_back(self, *a):
        self.manager.transition = SlideTransition(direction='right')
        self.manager.current = 'lab' if self.source == 'lab' else 'featured'


# ==================================================================
# گلچین برگزیده
# ==================================================================
class FeaturedScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='گلچین برگزیده', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(8))
        b_word = PillButton('خروجی Word', bg=C_BLUE, font_size='14sp')
        b_word.bind(on_release=lambda *a: self.export_word())
        b_clear = PillButton('پاک کردن کل', bg=C_RED, font_size='14sp')
        b_clear.bind(on_release=lambda *a: self.clear_all())
        top.add_widget(b_word)
        top.add_widget(b_clear)
        self.body(top)
        self.body(RLabel('برای دیدن نمونه‌های هر عملگر، روی آن بزنید.', font_size='13sp',
                         halign='center', color=C_MUTED, size_hint_y=None, height=dp(26)))
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.grid = GridLayout(cols=1, size_hint_y=None, spacing=dp(10), padding=dp(4))
        self.grid.bind(minimum_height=self.grid.setter('height'))
        self.scroll.add_widget(self.grid)
        self.body(self.scroll)

    def refresh(self):
        app = App.get_running_app()
        self.grid.clear_widgets()
        self.title_label.set_text('گلچین برگزیده (%d)' % len(app.featured))
        if not app.featured:
            self.grid.add_widget(RLabel('گلچین خالی است.', font_size='15sp', halign='center',
                                        color=C_MUTED, size_hint_y=None, height=dp(60)))
            return
        counts = {}
        for it in app.featured:
            k = op_of(it)
            counts[k] = counts.get(k, 0) + 1
        for key, title in OPERATORS:
            n = counts.get(key, 0)
            if n == 0:
                continue
            b = PillButton('%s  (%d نمونه)' % (title, n), bg=(0.16, 0.13, 0.05, 1),
                           size_hint_y=None, height=dp(64), font_size='16sp')
            b.bind(on_release=lambda inst, k=key, t=title: self.open_op(k, t))
            self.grid.add_widget(b)

    def open_op(self, key, title):
        scr = self.manager.get_screen('operator')
        scr.load('featured', key, title)
        self.manager.transition = SlideTransition(direction='left')
        self.manager.current = 'operator'

    def clear_all(self):
        app = App.get_running_app()

        def _do():
            app.featured = []
            app.save_featured()
            self.refresh()
        confirm('کل لیست گلچین پاک شود؟', _do, 'پاک کردن کل')

    def export_word(self):
        app = App.get_running_app()
        path = app.export_featured_word()
        if path:
            toast('فایل Word ساخته شد:\n' + path, 'خروجی Word')
        else:
            toast('گلچین خالی است یا خطایی رخ داد.', 'خطا')


# ==================================================================
# جستجو
# ==================================================================
class SearchScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='جستجو در کشفیات', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        self.q = PersianTextInput(hint_text=P('جستجو در لابراتوار و گلچین...'), font_size='15sp',
                                  background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1),
                                  padding=[dp(8), dp(14)])
        self.q.bind(on_text_validate=lambda *a: self.do_search())
        b = PillButton('جستجو', bg=C_BLUE, size_hint_x=None, width=dp(96), font_size='14sp')
        b.bind(on_release=lambda *a: self.do_search())
        top.add_widget(self.q)
        top.add_widget(b)
        self.body(top)
        self.info = RLabel('', font_size='13sp', halign='center', color=C_MUTED,
                           size_hint_y=None, height=dp(24))
        self.body(self.info)
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def refresh(self):
        self.list.clear_widgets()
        self.info.set_text('متن را وارد کنید و دکمهٔ جستجو را بزنید.')

    def _haystack(self, it):
        parts = [it.get('seed_arb', ''), it.get('target_arb', ''), it.get('seed_pers', ''),
                 it.get('target_pers', ''), it.get('mode', ''), it.get('relation_type', ''),
                 it.get('note', ''),
                 'سوره %s:%s' % (it.get('seed_s'), it.get('seed_a')),
                 'سوره %s:%s' % (it.get('target_s'), it.get('target_a'))]
        return core.strip_harakat(' '.join(str(x) for x in parts))

    def do_search(self):
        app = App.get_running_app()
        term = core.strip_harakat(self.q.query.strip())
        self.list.clear_widgets()
        if len(term) < 2:
            self.info.set_text('حداقل ۲ حرف وارد کنید.')
            return
        results = []
        for it in app.favs:
            if term in self._haystack(it):
                results.append(('lab', it))
        for it in app.featured:
            if term in self._haystack(it):
                results.append(('featured', it))
        self.info.set_text('%d نتیجه در لابراتوار و گلچین' % len(results))
        for source, it in results:
            self.list.add_widget(self._result_card(source, it))

    def _result_card(self, source, item):
        border = C_GOLD if source == 'featured' else C_BLUE
        tag = 'گلچین' if source == 'featured' else 'لابراتوار'
        card = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=border, orientation='vertical',
                        size_hint_y=None, padding=dp(10), spacing=dp(4))
        head = RLabel('[%s] %s' % (tag, item.get('mode', '')), bold=True, font_size='13sp',
                      color=(C_GOLD if source == 'featured' else C_BLUE), halign='right', size_hint_y=None)
        head.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(4)))
        card.add_widget(head)
        pair = RLabel('سوره %s:%s     سوره %s:%s' % (item.get('seed_s'), item.get('seed_a'),
                      item.get('target_s'), item.get('target_a')),
                      font_size='12sp', color=C_MUTED, halign='center', size_hint_y=None)
        pair.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(2)))
        card.add_widget(pair)
        a2 = RLabel('« %s »' % item.get('target_arb', ''), arabic=True, font_size='16sp',
                    halign='center', color=C_TEXT, size_hint_y=None)
        a2.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        card.add_widget(a2)
        btn = PillButton('مشاهدهٔ جزئیات', bg=C_BLUE, size_hint_y=None, height=dp(40), font_size='13sp')
        btn.bind(on_release=lambda *a, it=item, sc=source: show_discovery(it, sc, self))
        card.add_widget(btn)
        card.bind(minimum_height=lambda i, v: setattr(card, 'height', v + dp(20)))
        return card


# ==================================================================
# مدیریت برچسب‌ها
# ==================================================================
class TagsScreen(BaseScreen):
    DEFAULT = ["تقابل کامل", "گفت و گو", "زاویه دید متفاوت", "مکمل و بسط‌دهنده", "علت و معلول"]

    def __init__(self, **kw):
        super().__init__(title='مدیریت برچسب‌ها', **kw)
        top = BoxLayout(size_hint_y=None, height=dp(52), spacing=dp(8))
        self.q = PersianTextInput(hint_text=P('برچسب جدید...'), multiline=False, font_name='ui',
                           font_size='15sp', background_color=(1, 1, 1, 0.95),
                           foreground_color=(0.05, 0.08, 0.14, 1), padding=[dp(8), dp(14)])
        b = PillButton('افزودن', bg=C_GREEN, size_hint_x=None, width=dp(110), font_size='14sp')
        b.bind(on_release=lambda *a: self.add_tag())
        top.add_widget(self.q)
        top.add_widget(b)
        self.body(top)
        self.scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        self.list = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(8), padding=dp(4))
        self.list.bind(minimum_height=self.list.setter('height'))
        self.scroll.add_widget(self.list)
        self.body(self.scroll)

    def refresh(self):
        app = App.get_running_app()
        self.list.clear_widgets()
        for tag in app.get_all_tags():
            row = RoundBox(bg=(0.10, 0.14, 0.22, 1), orientation='horizontal', size_hint_y=None,
                           height=dp(50), padding=dp(10), spacing=dp(8))
            row.add_widget(RLabel(tag, font_size='15sp', halign='right'))
            if tag not in self.DEFAULT and tag != 'نامشخص':
                b = PillButton('', bg=C_RED, size_hint_x=None, width=dp(56), font_size='14sp')
                b.bind(on_release=lambda x, t=tag: self.del_tag(t))
                row.add_widget(b)
            self.list.add_widget(row)

    def add_tag(self):
        app = App.get_running_app()
        t = self.q.query.strip()
        if not t:
            return
        if t in app.get_all_tags():
            toast('این برچسب قبلاً وجود دارد.', 'تکرار')
            return
        app.user_tags.append(t)
        app.save_user_tags()
        self.q.clear_logical()
        self.refresh()

    def del_tag(self, tag):
        app = App.get_running_app()
        def _do():
            if tag in app.user_tags:
                app.user_tags.remove(tag)
            for it in app.favs:
                if it.get('relation_type') == tag:
                    it['relation_type'] = 'نامشخص'
            app.save_user_tags()
            app.save_favs()
            self.refresh()
        confirm(f'برچسب «{tag}» حذف شود؟ (کشفیات مربوطه به نامشخص تغییر می‌کند)', _do, 'حذف برچسب')


# ==================================================================
# رسانه و معرفی
# ==================================================================
class MediaScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='رسانه و معرفی', **kw)
        self.sound = None
        self.video = None
        scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(12), padding=dp(6))
        box.bind(minimum_height=box.setter('height'))

        box.add_widget(RLabel('ویدیوی معرفی', bold=True, font_size='18sp', color=C_GOLD,
                              halign='center', size_hint_y=None, height=dp(34)))
        self.video_holder = RoundBox(bg=(0, 0, 0, 0.5), orientation='vertical', size_hint_y=None,
                                     height=dp(220), padding=dp(6))
        box.add_widget(self.video_holder)
        note = RLabel('برای نمایش مناسب در موبایل، کلیپ خود را با نام intro.mp4 در پوشهٔ assets قرار دهید.',
                     font_size='12sp', halign='center', color=C_MUTED, size_hint_y=None, height=dp(44))
        box.add_widget(note)

        box.add_widget(RLabel('صدا', bold=True, font_size='18sp', color=C_GOLD,
                              halign='center', size_hint_y=None, height=dp(34)))
        b_designer = PillButton('چند کلام از طراح در مورد اپلیکیشن', bg=C_ORANGE, size_hint_y=None, height=dp(52), font_size='14sp')
        b_designer.bind(on_release=lambda *a: self.play('designer.mp3'))
        box.add_widget(b_designer)
        b_stop = PillButton('توقف', bg=C_RED, size_hint_y=None, height=dp(52))
        b_stop.bind(on_release=lambda *a: self.stop())
        box.add_widget(b_stop)
        box.add_widget(Widget(size_hint_y=None, height=dp(20)))

        scroll.add_widget(box)
        self.body(scroll)

    def refresh(self):
        # بارگذاری ویدیو در صورت وجود
        self.video_holder.clear_widgets()
        vpath = asset('intro.mp4')
        if not os.path.exists(vpath):
            self.video_holder.add_widget(RLabel('کلیپ intro.mp4 یافت نشد.\nفایل خود را در assets بگذارید.',
                                                halign='center', color=C_MUTED, font_size='14sp'))
            return
        try:
            from kivy.uix.video import Video
        except Exception:
            self.video_holder.add_widget(RLabel('پخش ویدیو در دسترس نیست.', halign='center', color=C_MUTED))
            return
        try:
            self.video = Video(source=vpath, state='stop', options={'eos': 'stop'},
                               allow_stretch=True, keep_ratio=True, size_hint_y=1)
            self.video.bind(on_load=lambda *a: None)
            self.video_holder.add_widget(self.video)
            self.ctrl_btn = PillButton('پخش ویدیو', bg=C_BLUE, size_hint_y=None, height=dp(42),
                                       font_size='13sp')
            self.ctrl_btn.bind(on_release=lambda *a: self.toggle_video())
            self.video_holder.add_widget(self.ctrl_btn)
        except Exception as e:
            print('video error:', e)
            self.video_holder.add_widget(RLabel('پخش این ویدیو ممکن نشد. برای پخش mp4 روی ویندوز، بسته ffpyplayer نصب شود:\npy -3.12 -m pip install ffpyplayer',
                                                halign='center', color=C_MUTED, font_size='12sp'))

    def toggle_video(self):
        if not self.video:
            return
        if self.video.state == 'play':
            self.video.state = 'pause'
            if hasattr(self, 'ctrl_btn'):
                self.ctrl_btn.set_text('پخش ویدیو')
        else:
            try:
                if self.video.eos or (self.video.position and self.video.duration
                                      and self.video.position >= self.video.duration - 0.2):
                    self.video.position = 0
            except Exception:
                pass
            self.video.state = 'play'
            self.video.volume = 1
            if hasattr(self, 'ctrl_btn'):
                self.ctrl_btn.set_text('توقف ویدیو')

    def play(self, name):
        from kivy.core.audio import SoundLoader
        self.stop()
        path = asset(name)
        if not os.path.exists(path):
            toast('فایل صوتی یافت نشد.', 'خطا')
            return
        self.sound = SoundLoader.load(path)
        if self.sound:
            self.sound.play()

    def stop(self):
        if self.sound:
            self.sound.stop()
            self.sound = None

    def go_back(self, *a):
        self.stop()
        if self.video:
            self.video.state = 'stop'
        super().go_back()


# ==================================================================
# پشتیبان و بازیابی
# ==================================================================
class BackupScreen(BaseScreen):
    def __init__(self, **kw):
        super().__init__(title='پشتیبان و بازیابی', **kw)
        box = BoxLayout(orientation='vertical', spacing=dp(14), padding=dp(16))
        box.add_widget(RLabel('از کشفیات، گلچین و برچسب‌های خود نسخهٔ پشتیبان بگیرید یا آن را بازیابی کنید.',
                              font_size='15sp', halign='center', color=C_MUTED, size_hint_y=None, height=dp(70)))
        b_backup = PillButton('ساخت فایل پشتیبان (zip)', bg=C_GREEN, size_hint_y=None, height=dp(56))
        b_backup.bind(on_release=lambda *a: self.backup())
        box.add_widget(b_backup)
        self.info = RLabel('', font_size='13sp', halign='center', color=C_GOLD, size_hint_y=None, height=dp(60))
        box.add_widget(self.info)
        box.add_widget(Widget())
        self.body(box)

    def refresh(self):
        self.info.set_text('')

    def backup(self):
        # ابتدا تلاش برای پنجرهٔ انتخاب پوشه (ویندوز/دسکتاپ)
        chosen = None
        try:
            import tkinter
            from tkinter import filedialog
            root = tkinter.Tk()
            root.withdraw()
            root.attributes('-topmost', True)
            chosen = filedialog.askdirectory(title='پوشهٔ ذخیرهٔ پشتیبان را انتخاب کنید')
            root.destroy()
            if not chosen:
                self.info.set_text('ذخیره‌سازی لغو شد.')
                return
            self._do_backup(chosen)
        except Exception:
            # در موبایل/محیطی که پنجرهٔ انتخاب پوشه ندارد، مسیر را بپرس
            self._ask_path_popup()

    def _ask_path_popup(self):
        content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(10))
        content.add_widget(RLabel('مسیر (پوشهٔ) ذخیرهٔ پشتیبان را وارد کنید:',
                                  font_size='14sp', size_hint_y=None, height=dp(30)))
        app = App.get_running_app()
        ti = PersianTextInput(multiline=False, font_size='14sp', size_hint_y=None, height=dp(48),
                              background_color=(1, 1, 1, 0.95), foreground_color=(0.05, 0.08, 0.14, 1))
        ti.set_logical(app.store_dir)
        content.add_widget(ti)
        content.add_widget(RLabel('خالی بگذارید تا در پوشهٔ پیش‌فرض برنامه ذخیره شود.',
                                  font_size='12sp', color=C_MUTED, size_hint_y=None, height=dp(28)))
        p = Popup(title=P('محل ذخیرهٔ پشتیبان'), content=content, size_hint=(0.92, 0.5),
                  title_font='ui', title_align='center', separator_color=C_GOLD)
        row = BoxLayout(size_hint_y=None, height=dp(48), spacing=dp(10))
        ok = PillButton('ذخیره', bg=C_GREEN)

        def _ok(*a):
            d = ti.query.strip()
            p.dismiss()
            self._do_backup(d if d else None)
        ok.bind(on_release=_ok)
        cn = PillButton('انصراف', bg=C_RED)
        cn.bind(on_release=p.dismiss)
        row.add_widget(ok)
        row.add_widget(cn)
        content.add_widget(row)
        p.open()

    def _do_backup(self, dest_dir):
        app = App.get_running_app()
        if dest_dir and not os.path.isdir(dest_dir):
            self.info.set_text('مسیر نامعتبر است. دوباره تلاش کنید.')
            return
        try:
            path = app.make_backup(dest_dir)
            self.info.set_text('فایل پشتیبان ساخته شد:\n' + path)
            toast('پشتیبان با موفقیت ذخیره شد.', 'پشتیبان')
        except Exception as e:
            self.info.set_text('خطا در ساخت پشتیبان:\n' + str(e))


# ==================================================================
# درباره
# ==================================================================
class AboutScreen(BaseScreen):
    WEBSITE = 'https://6a304b9599e34.mywebzi.ir/'
    BALE_URL = 'https://ble.ir/dr_parsa114'

    def __init__(self, **kw):
        super().__init__(title='درباره', **kw)
        scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        box = BoxLayout(orientation='vertical', spacing=dp(12), padding=dp(16), size_hint_y=None)
        box.bind(minimum_height=box.setter('height'))
        box.add_widget(self._lbl('قطب‌نمای قرآنی', bold=True, font_size='24sp', color=C_GOLD, halign='center'))
        box.add_widget(self._lbl('پردازش آینه‌ای (هولوگرافیک) — نسخهٔ موبایل', font_size='15sp', color=C_TEXT, halign='center'))
        box.add_widget(self._lbl('کاوش الگوهای عددی و معنایی میان آیات قرآن کریم. تمام ۶۲۳۶ آیه به صورت آفلاین در اپ گنجانده شده است.', font_size='13sp', color=C_MUTED, halign='center'))
        box.add_widget(Widget(size_hint_y=None, height=dp(8)))
        box.add_widget(self._lbl('راه ارتباطی با مؤلف:', bold=True, font_size='17sp', color=C_GOLD, halign='right'))
        b_site = PillButton('سایت مرجع قرآن ابر ماتریس', bg=C_BLUE, size_hint_y=None, height=dp(56), font_size='15sp')
        b_site.bind(on_release=lambda *a: self.open_url(self.WEBSITE))
        box.add_widget(b_site)
        box.add_widget(self._lbl(self.WEBSITE, font_size='12sp', color=C_MUTED, halign='center'))
        b_bale = PillButton('ارتباط در پیام‌رسان بله:  dr_parsa114', bg=C_GREEN, size_hint_y=None, height=dp(56), font_size='15sp')
        b_bale.bind(on_release=lambda *a: self.open_url(self.BALE_URL))
        box.add_widget(b_bale)
        box.add_widget(Widget(size_hint_y=None, height=dp(8)))
        box.add_widget(self._card_text('لطفاً کشفیات ویژهٔ خود را با مؤلف در میان بگذارید و به اشتراک بگذارید تا در نسخه‌های بعدی گنجانده شود.'))
        box.add_widget(self._card_text('این اپلیکیشن و سامانهٔ پردازش آن در حال توسعه و تکامل است؛ ان‌شاءالله به لطف خالق هستی و با کمک یکدیگر، با بزرگ‌تر کردن فهرست آیات آینه‌ای به این هدف مهم دست خواهیم یافت.'))
        box.add_widget(Widget(size_hint_y=None, height=dp(16)))
        scroll.add_widget(box)
        self.body(scroll)

    def _lbl(self, text, **kw):
        kw.setdefault('size_hint_y', None)
        l = RLabel(text, **kw)
        l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(12)))
        return l

    def _card_text(self, text):
        c = RoundBox(bg=(0.10, 0.14, 0.22, 1), border=C_GOLD, orientation='vertical', size_hint_y=None, padding=dp(12))
        l = RLabel(text, font_size='14sp', color=C_TEXT, halign='right', size_hint_y=None)
        l.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(6)))
        c.add_widget(l)
        c.bind(minimum_height=lambda i, v: setattr(c, 'height', v + dp(24)))
        return c

    def open_url(self, url):
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception:
            toast('نشانی: ' + url, 'لینک')

    def refresh(self):
        pass


# ==================================================================
# راهنما
# ==================================================================
class GuideScreen(BaseScreen):
    SECTIONS = [
        ('انتخاب بذر (سوره و آیه)', 'در صفحهٔ اصلی شمارهٔ سوره و آیه را وارد کنید. این «بذر» مبنای همهٔ پردازش‌هاست. اگر آیه‌ای خارج از محدوده وارد شود، به نزدیک‌ترین آیهٔ معتبر اصلاح می‌شود.'),
        ('پردازش ماتریس', 'هفت عملگر آینه‌ای (جابجایی و تقارن سوره/آیه) را روی بذر اعمال می‌کند و هفت آیهٔ مقصد را با متن کامل عربی و ترجمه نشان می‌دهد. هر مقصد را با دکمهٔ «ثبت این کشف» در لابراتوار ذخیره کنید.'),
        ('پیش‌بینی (معنا)', 'مقاصد آینه‌ای را بر اساس اشتراک واژگانی و ارتباط معنایی با بذر رتبه‌بندی می‌کند تا محتمل‌ترین تناظرها بالاتر بیایند.'),
        ('پیش‌بینی (اعداد)', 'با فیلترهای عددی مانند نیم‌کرهٔ سوره، اثر انگشت رقمی و تلورانس، نامزدهای عددی را غربال و اولویت‌بندی می‌کند.'),
        ('لابراتوار کشفیات', 'همهٔ کشف‌های ثبت‌شدهٔ شما اینجاست و بر اساس هفت عملگر دسته‌بندی شده. روی هر عملگر بزنید تا کشف‌های همان دسته باز شود؛ سپس روی هر کشف بزنید تا جزئیات کامل (عربی + ترجمهٔ مبدأ و مقصد) با امکان گلچین، ویرایش تحلیل، حذف و کپی را ببینید.'),
        ('گلچین برگزیده', 'کشف‌های مهم را از لابراتوار به گلچین می‌آورید. اینجا هم مانند لابراتوار بر اساس عملگرها دسته‌بندی شده و می‌توانید از کل گلچین خروجی Word بگیرید.'),
        ('جستجوی آیات', 'در میان کشفیات لابراتوار و گلچین جستجو می‌کند (نه کل قرآن). متن عربی، ترجمه، برچسب و تحلیل شما جستجو می‌شود.'),
        ('مدیریت برچسب‌ها', 'برچسب‌های «رفتار آیه» (مانند تقابل کامل، گفت‌وگو، علت و معلول) را می‌سازید یا حذف می‌کنید تا هنگام ثبت تحلیل به کشف‌ها نسبت دهید.'),
        ('رسانه و معرفی', 'ویدیوی معرفی (intro.mp4) و صدای طراح را اینجا می‌بینید و می‌شنوید. صدای کوتاه معرفی هنگام باز شدن برنامه یک‌بار پخش می‌شود.'),
        ('پشتیبان و بازیابی', 'از داده‌های خود (کشفیات، گلچین، برچسب‌ها) نسخهٔ پشتیبان zip بگیرید تا اطلاعاتتان امن بماند.'),
        ('درباره', 'معرفی اپلیکیشن و راه‌های ارتباط با مؤلف (سایت مرجع و شناسهٔ پیام‌رسان بله).'),
    ]

    def __init__(self, **kw):
        super().__init__(title='راهنما', **kw)
        scroll = ScrollView(do_scroll_x=False, bar_width=dp(4))
        box = BoxLayout(orientation='vertical', size_hint_y=None, spacing=dp(10), padding=dp(8))
        box.bind(minimum_height=box.setter('height'))
        box.add_widget(RLabel('روی هر بخش بزنید تا توضیح کامل آن باز شود.', font_size='14sp',
                              color=C_MUTED, halign='center', size_hint_y=None, height=dp(34)))
        for title, body in self.SECTIONS:
            b = PillButton(title, bg=(0.16, 0.13, 0.05, 1), size_hint_y=None, height=dp(56), font_size='15sp')
            b.bind(on_release=lambda inst, t=title, d=body: self.show_help(t, d))
            box.add_widget(b)
        gt = asset('guide_table.png')
        if os.path.exists(gt):
            img = Image(source=gt, size_hint_y=None, height=dp(200), allow_stretch=True, keep_ratio=True)
            box.add_widget(img)
        scroll.add_widget(box)
        self.body(scroll)

    def show_help(self, title, body):
        content = BoxLayout(orientation='vertical', padding=dp(14), spacing=dp(10))
        sc = ScrollView(do_scroll_x=False, bar_width=dp(4))
        inner = BoxLayout(orientation='vertical', size_hint_y=None, padding=dp(4))
        inner.bind(minimum_height=inner.setter('height'))
        lbl = RLabel(body, font_size='15sp', color=C_TEXT, halign='right', size_hint_y=None)
        lbl.bind(texture_size=lambda i, v: setattr(i, 'height', v[1] + dp(10)))
        inner.add_widget(lbl)
        sc.add_widget(inner)
        content.add_widget(sc)
        p = Popup(title=P(title), content=content, size_hint=(0.92, 0.6),
                  title_font='ui', title_align='center', separator_color=C_GOLD)
        btn = PillButton('بستن', bg=C_BLUE, size_hint_y=None, height=dp(46))
        btn.bind(on_release=p.dismiss)
        content.add_widget(btn)
        p.open()

    def refresh(self):
        pass


# ==================================================================
# اپلیکیشن
# ==================================================================
class QuranMirrorApp(App):
    def build(self):
        self.title = 'قطب‌نمای قرآنی'
        Window.clearcolor = C_BG
        # داده
        self.data = core.QuranData(asset('datakavosh.csv'))
        self._init_storage()
        self.load_favs()
        self.load_featured()
        self.load_user_tags()
        # صفحات
        sm = ScreenManager(transition=FadeTransition(duration=0.25))
        sm.add_widget(HomeScreen(name='home'))
        sm.add_widget(MatrixScreen(name='matrix'))
        sm.add_widget(PredictScreen(name='predict'))
        sm.add_widget(LabScreen(name='lab'))
        sm.add_widget(OperatorScreen(name='operator'))
        sm.add_widget(FeaturedScreen(name='featured'))
        sm.add_widget(SearchScreen(name='search'))
        sm.add_widget(TagsScreen(name='tags'))
        sm.add_widget(MediaScreen(name='media'))
        sm.add_widget(BackupScreen(name='backup'))
        sm.add_widget(AboutScreen(name='about'))
        sm.add_widget(GuideScreen(name='guide'))
        self.sm = sm
        return sm

    def on_start(self):
        try:
            from kivy.core.audio import SoundLoader
            path = asset('voice.mp3')
            if os.path.exists(path):
                snd = SoundLoader.load(path)
                if snd:
                    self._intro_sound = snd
                    Clock.schedule_once(lambda *a: snd.play(), 0.6)
        except Exception:
            pass

    # ---------- ذخیره‌سازی ----------
    def _init_storage(self):
        self.store_dir = self.user_data_dir
        for name in ('favorites.json', 'featured.json', 'user_tags.json'):
            dst = os.path.join(self.store_dir, name)
            if not os.path.exists(dst):
                src = asset(name)
                try:
                    if os.path.exists(src):
                        shutil.copy(src, dst)
                    else:
                        with open(dst, 'w', encoding='utf-8') as f:
                            json.dump([], f, ensure_ascii=False)
                except Exception:
                    pass

    def _p(self, name):
        return os.path.join(self.store_dir, name)

    def load_favs(self):
        try:
            with open(self._p('favorites.json'), encoding='utf-8') as f:
                self.favs = json.load(f)
        except Exception:
            self.favs = []

    def save_favs(self):
        with open(self._p('favorites.json'), 'w', encoding='utf-8') as f:
            json.dump(self.favs, f, ensure_ascii=False, indent=2)

    def load_featured(self):
        try:
            with open(self._p('featured.json'), encoding='utf-8') as f:
                self.featured = json.load(f)
        except Exception:
            self.featured = []

    def save_featured(self):
        with open(self._p('featured.json'), 'w', encoding='utf-8') as f:
            json.dump(self.featured, f, ensure_ascii=False, indent=2)

    def load_user_tags(self):
        try:
            with open(self._p('user_tags.json'), encoding='utf-8') as f:
                self.user_tags = json.load(f)
        except Exception:
            self.user_tags = []

    def save_user_tags(self):
        with open(self._p('user_tags.json'), 'w', encoding='utf-8') as f:
            json.dump(self.user_tags, f, ensure_ascii=False, indent=2)

    def get_all_tags(self):
        tags = set(TagsScreen.DEFAULT) | {'نامشخص'}
        for it in self.favs:
            t = it.get('relation_type')
            if t:
                tags.add(t)
        tags.update(self.user_tags)
        return sorted(tags)

    # ---------- عملیات کشف ----------
    def add_discovery(self, seed, target):
        entry = {
            'mode': target.get('mode', ''),
            'seed_s': seed['s'], 'seed_a': seed['a'],
            'seed_arb': seed['arb'], 'seed_pers': seed['pers'],
            'target_s': target['s'], 'target_a': target['a'],
            'target_arb': target.get('arb', ''), 'target_pers': target.get('pers', ''),
            'note': '', 'relation_type': 'نامشخص',
            'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        }
        # جلوگیری از تکرار
        for it in self.favs:
            if (it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'),
                    it.get('mode')) == (entry['seed_s'], entry['seed_a'], entry['target_s'],
                                        entry['target_a'], entry['mode']):
                toast('این کشف قبلاً ثبت شده است.', 'تکرار')
                return
        self.favs.append(entry)
        self.save_favs()
        toast('کشف در لابراتوار ثبت شد. ', 'ثبت کشف')

    def add_featured(self, item, screen=None):
        for it in self.featured:
            if (it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a')) == \
               (item.get('seed_s'), item.get('seed_a'), item.get('target_s'), item.get('target_a')):
                toast('این مورد در گلچین هست.', 'تکرار')
                return
        self.featured.append(dict(item))
        self.save_featured()
        toast('به گلچین اضافه شد. ', 'گلچین')

    def add_all_featured(self):
        existing = {(it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'))
                    for it in self.featured}
        n = 0
        for it in self.favs:
            key = (it.get('seed_s'), it.get('seed_a'), it.get('target_s'), it.get('target_a'))
            if key not in existing:
                self.featured.append(dict(it))
                existing.add(key)
                n += 1
        self.save_featured()
        return n

    # ---------- خروجی Word ----------
    def export_featured_word(self):
        if not self.featured:
            return None
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()
            h = doc.add_heading('', level=0)
            run = h.add_run('گلچین آیات آینه‌ای')
            run.font.size = Pt(20)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, it in enumerate(self.featured, 1):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(f"{i}. [{it.get('mode', '')}] سوره {it.get('seed_s')}:{it.get('seed_a')} سوره {it.get('target_s')}:{it.get('target_a')}")
                r.bold = True
                for key in ('seed_arb', 'seed_pers', 'target_arb', 'target_pers'):
                    par = doc.add_paragraph(it.get(key, ''))
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if it.get('note'):
                    par = doc.add_paragraph('یادداشت: ' + it['note'])
                    par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                doc.add_paragraph('—' * 20)
            out = self._p('golchin_%s.docx' % datetime.now().strftime('%Y%m%d_%H%M'))
            doc.save(out)
            return out
        except Exception as e:
            print('word export error:', e)
            return None

    # ---------- پشتیبان ----------
    def make_backup(self, dest_dir=None):
        fname = 'backup_%s.zip' % datetime.now().strftime('%Y%m%d_%H%M')
        if dest_dir:
            out = os.path.join(dest_dir, fname)
        else:
            out = self._p(fname)
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
            for name in ('favorites.json', 'featured.json', 'user_tags.json'):
                pth = self._p(name)
                if os.path.exists(pth):
                    z.write(pth, name)
        return out


if __name__ == '__main__':
    QuranMirrorApp().run()
